"""
한글 (HWP / HWPX / PDF / PPTX) ➔ Word (.docx) 변환 모듈

파이프라인 아키텍처:
1. [1순위: 시각적 재현 트랙 (인쇄·열람·제출용)]
   rhwp 고화질 PDF 렌더링 ➔ pdf2docx 레이아웃/스타일/위상 역공학
   - 장점: 원본 문서의 4쪽 분량, 글꼴 크기 위계, 행 높이, 여백, 박스 배치 완벽 일치
   - 주의: PDF 좌표 기반이므로 띄어쓰기 소실 및 라벨-값 분리 현상 발생 (RAG/학습 데이터로는 사용 금지)

2. [2순위: 구조 보존 폴백 트랙 (HWPX XML 네이티브 테이블 생성)]
   HWPX XML 셀 좌표(cellAddr) + 병합(cellSpan) + 열너비(cellSz) ➔ python-docx 네이티브 표 생성
   - 장점: 텍스트 무손실, 병합 격자 유지, 편집 용이
   - pdf2docx 실패 시 완벽한 폴백으로 작동

3. [3순위: 평문 폴백 트랙]
   Unified Parser 마크다운 ➔ python-docx 기본 문단 생성
"""
import logging
import warnings
from pathlib import Path

# Suppress PyMuPDF deprecation warnings from third party libraries
warnings.filterwarnings("ignore", category=UserWarning)

logger = logging.getLogger(__name__)

# HWPUNIT: 1/7200 inch
HWPUNIT_PER_INCH = 7200


def _clear_cell(cell) -> None:
    """셀을 빈 문단 하나만 남긴 상태로 정리합니다."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for run in list(cell.paragraphs[0].runs):
        run._element.getparent().remove(run._element)


def _apply_table_borders(table) -> None:
    """표 전체에 실선 테두리를 적용합니다."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    edges = "".join(
        f'<w:{e} w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        for e in ("top", "left", "bottom", "right", "insideH", "insideV")
    )
    table._tbl.tblPr.append(parse_xml(f'<w:tblBorders {nsdecls("w")}>{edges}</w:tblBorders>'))


def _set_fixed_layout(table) -> None:
    """열 너비를 지정한 대로 고정합니다."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    table._tbl.tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))


def _shade_cell(cell, hex_color: str) -> None:
    """셀에 배경색을 적용합니다."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    )


def _render_table(doc, table_data: dict) -> None:
    """셀 좌표·병합·열너비를 보존한 채 네이티브 Word 표를 생성합니다."""
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    row_count = table_data.get("row_count") or 0
    col_count = table_data.get("col_count") or 0
    cells = table_data.get("cells") or []
    if row_count <= 0 or col_count <= 0 or not cells:
        return

    table = doc.add_table(rows=row_count, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _apply_table_borders(table)
    _set_fixed_layout(table)

    # 1. 셀 병합 수행
    for c in cells:
        if c["col_span"] <= 1 and c["row_span"] <= 1:
            continue
        r0, c0 = c["row"], c["col"]
        r1 = min(r0 + c["row_span"] - 1, row_count - 1)
        c1 = min(c0 + c["col_span"] - 1, col_count - 1)
        if r1 <= r0 and c1 <= c0:
            continue
        try:
            table.cell(r0, c0).merge(table.cell(r1, c1))
        except Exception as e:
            logger.warning(f"셀 병합 건너뜀 (r{r0},c{c0} ➔ r{r1},c{c1}): {e}")

    # 2. 텍스트와 배경색 채우기
    for c in cells:
        try:
            cell = table.cell(c["row"], c["col"])
        except Exception:
            continue
        _clear_cell(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if c.get("fill"):
            _shade_cell(cell, c["fill"])

        if c.get("text"):
            para = cell.paragraphs[0]
            run = para.add_run(c["text"])
            run.font.size = Pt(9.5)
            if c.get("fill"):
                run.font.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. 열 너비 반영
    for idx, width in enumerate(table_data.get("col_widths") or []):
        if not width or idx >= col_count:
            continue
        inches = Inches(width / HWPUNIT_PER_INCH)
        for row in table.rows:
            try:
                row.cells[idx].width = inches
            except Exception:
                pass

    doc.add_paragraph().paragraph_format.space_before = Pt(4)


def render_blocks_to_docx(blocks: list, output_docx_path: str | Path) -> Path:
    """구조화 블록 리스트를 네이티브 Word 문서로 렌더링합니다."""
    import docx
    from docx.shared import Pt, Inches

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    for block in blocks:
        if block.get("type") == "table" and block.get("table"):
            _render_table(doc, block["table"])
            continue

        text = (block.get("text") or "").strip()
        if not text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2
        p.add_run(text)

    out_path = Path(output_docx_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def convert_hwp_to_docx(input_path: str | Path, output_dir: str | Path | None = None) -> Path:
    """Convert a Hangul (.hwp, .hwpx), PDF, or Office document to a high-fidelity Word (.docx) file.

    Uses 1st priority pdf2docx for visual reconstruction, with 2nd priority HWPX structure fallback.

    Args:
        input_path: Path to source document (.hwp, .hwpx, .pdf, .pptx, .md, .txt).
        output_dir: Target directory. Defaults to input file's parent folder.

    Returns:
        Path to the generated .docx file.
    """
    input_path = Path(input_path)
    if not input_path.is_file():
        raise FileNotFoundError(f"입력 파일을 찾을 수 없습니다: {input_path}")

    if output_dir is None:
        target_dir = input_path.parent
    else:
        target_dir = Path(output_dir)
        target_dir.mkdir(parents=True, exist_ok=True)

    output_docx_path = target_dir / (input_path.stem + ".docx")
    ext = input_path.suffix.lower()

    # 1. [1순위 최우선: 시각 재현 트랙] rhwp 고화질 PDF ➔ pdf2docx 정밀 레이아웃/표 복원
    temp_pdf_dir = None
    try:
        from pdf2docx import Converter
        from src.utils.hwp_to_pdf import convert_hwp_to_pdf

        if ext == ".pdf":
            temp_pdf_path = input_path
            need_cleanup = False
        else:
            temp_pdf_dir = target_dir / "_temp_pdf"
            temp_pdf_dir.mkdir(parents=True, exist_ok=True)
            temp_pdf_path = convert_hwp_to_pdf(input_path, output_dir=temp_pdf_dir)
            need_cleanup = True

        cv = Converter(str(temp_pdf_path))
        cv.convert(str(output_docx_path))
        cv.close()

        if need_cleanup and temp_pdf_path.is_file():
            try:
                temp_pdf_path.unlink()
                if temp_pdf_dir and temp_pdf_dir.is_dir() and not any(temp_pdf_dir.iterdir()):
                    temp_pdf_dir.rmdir()
            except Exception:
                pass

        if output_docx_path.is_file() and output_docx_path.stat().st_size > 0:
            logger.info(f"[pdf2docx] 최고품질 시각 레이아웃 Word 변환 성공: {output_docx_path.name}")
            return output_docx_path

    except Exception as e:
        logger.warning(f"[pdf2docx] 레이아웃 복원 실패 ({e}). HWPX 구조 보존 트랙으로 폴백합니다.")

    # 2. [2순위 폴백: 구조 보존 트랙] HWPX XML 구조 파싱 ➔ 네이티브 Word 테이블 복원
    if ext in (".hwp", ".hwpx"):
        try:
            from src.parsers.unified_parser import UnifiedDocumentParser

            parsed = UnifiedDocumentParser().parse(input_path)
            blocks = parsed.get("blocks") or []
            table_blocks = [b for b in blocks if b.get("type") == "table"]

            if table_blocks:
                render_blocks_to_docx(blocks, output_docx_path)
                if output_docx_path.is_file() and output_docx_path.stat().st_size > 0:
                    logger.info(f"[structured] 구조 보존 폴백 Word 생성 성공: {output_docx_path.name} (표 {len(table_blocks)}개)")
                    return output_docx_path
        except Exception as e:
            logger.warning(f"[structured] 구조 보존 폴백 실패 ({e}). 평문 빌더로 전환합니다.")

    # 3. [3순위 폴백: 평문 트랙] 마크다운 ➔ python-docx
    try:
        import docx
        from src.parsers.unified_parser import UnifiedDocumentParser

        parsed_res = UnifiedDocumentParser().parse(input_path)
        markdown_text = parsed_res.get("markdown", "")

        doc = docx.Document()
        for line in markdown_text.splitlines():
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("* ") or line.startswith("- "):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(str(output_docx_path))
        if output_docx_path.is_file() and output_docx_path.stat().st_size > 0:
            return output_docx_path

    except Exception as e:
        logger.error(f"[Fallback] 마크다운 docx 생성도 실패: {e}")

    raise RuntimeError(f"Word(.docx) 변환 실패: {output_docx_path.name} 생성 불가")


# Universal alias
convert_document_to_docx = convert_hwp_to_docx
convert_to_docx = convert_hwp_to_docx
